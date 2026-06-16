#!/usr/bin/env python3
"""
施工日志自动生成

从 Excel 进度表读取数据，自动填充 Word 模板，生成每日施工日志。

使用示例：
  python daily_log_gen.py --template 日志模板.docx --data 进度表.xlsx --date 2026-06-15
"""

import argparse
from pathlib import Path
from datetime import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt


def generate_logs(template_file: str, data_file: str,
                  date_str: str, output: str):
    """从 Excel 数据生成施工日志"""
    template_path = Path(template_file)
    data_path = Path(data_file)

    if not template_path.exists():
        print(f"⚠️  模板文件不存在: {template_file}")
        print("   将使用内置默认模板生成")
        use_template = False
    else:
        use_template = True

    if not data_path.exists():
        print(f"❌ 数据文件不存在: {data_file}")
        return

    # 读取施工数据
    df = pd.read_excel(data_path) if data_path.suffix in (".xlsx", ".xls") \
         else pd.read_csv(data_path)

    print(f"📂 读取施工数据: {len(df)} 条记录")

    target_date = datetime.strptime(date_str, "%Y-%m-%d")
    day_data = df.copy()  # 简化：使用全部数据。实际可加日期筛选

    # 生成日志
    if use_template:
        doc = Document(template_path)
    else:
        doc = Document()
        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(11)

    # ── 替换占位符 ──
    date_cn = target_date.strftime("%Y年%m月%d日")
    weekday = ["一", "二", "三", "四", "五", "六", "日"][target_date.weekday()]

    placeholders = {
        "{{日期}}": date_cn,
        "{{星期}}": f"星期{weekday}",
        "{{天气}}": "晴",
        "{{温度}}": "15°C ~ 28°C",
        "{{施工内容}}": _build_content(day_data),
        "{{施工人员}}": f"管理人员 3 人，作业人员 12 人",
        "{{机械设备}}": _build_equipment(day_data),
        "{{材料进场}}": _build_materials(day_data),
        "{{安全问题}}": "无",
        "{{备注}}": "",
    }

    # 替换段落中的占位符
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, str(value))

    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    doc.save(output)
    print(f"✅ 施工日志已生成: {output}")


def _build_content(df):
    """从数据构建施工内容描述"""
    lines = []
    if "工序" in df.columns:
        for _, row in df.head(10).iterrows():
            location = row.get("部位", row.get("位置", ""))
            task = row.get("工序", row.get("工作内容", ""))
            qty = row.get("数量", "")
            unit = row.get("单位", "")
            qty_str = f"（{qty}{unit}）" if qty and str(qty) != "nan" else ""
            lines.append(f"  • {location}：{task}{qty_str}")
    else:
        lines.append("  • 按施工计划正常进行")
    return "\n".join(lines) if lines else "按施工计划正常进行"


def _build_equipment(df):
    """提取机械设备信息"""
    if "设备" in df.columns:
        equip = df["设备"].dropna().unique()
        return "、".join(equip[:5])
    return "汽车吊 1 台，挖掘机 1 台"


def _build_materials(df):
    """提取材料信息"""
    if "材料" in df.columns:
        mats = df["材料"].dropna().unique()
        return "、".join(mats[:5])
    return "按计划进场"


def main():
    parser = argparse.ArgumentParser(
        description="施工日志自动生成 — Excel 数据 + Word 模板 → 日志文件"
    )
    parser.add_argument("--template", default=None,
                        help="Word 模板路径（可选，不含模板则用默认格式）")
    parser.add_argument("--data", required=True,
                        help="施工进度数据 Excel/CSV 文件")
    parser.add_argument("--date", required=True,
                        help="日志日期，格式 YYYY-MM-DD")
    parser.add_argument("--output", default=None,
                        help="输出文件路径（默认: 施工日志_日期.docx）")
    args = parser.parse_args()

    if args.output is None:
        args.output = f"施工日志_{args.date}.docx"

    generate_logs(args.template, args.data, args.date, args.output)


if __name__ == "__main__":
    main()
